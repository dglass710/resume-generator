# generator.py
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
import os
import sys
from pathlib import Path
import re

# Refined token checks using regular expressions.
_email_re = re.compile(r'^[\w\.-]+@[\w\.-]+\.\w+$')
_phone_re = re.compile(r'^\D*(\d\D*){10}$')
_url_re = re.compile(r'^(https?://|www\.)', re.IGNORECASE)

def set_single_spacing(paragraph):
    """Set single spacing and remove extra space before/after the paragraph."""
    p_format = paragraph.paragraph_format
    p_format.line_spacing = 1
    p_format.space_after = Pt(0)
    p_format.space_before = Pt(0)

def add_hyperlink(paragraph, url, text):
    """Inserts a hyperlink into a paragraph. Returns the hyperlink element."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    # Create the w:hyperlink element and set its relationship id.
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create a run inside the hyperlink.
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Set hyperlink styling (black, underlined) to match Harvard template
    color = OxmlElement('w:color')
    color.set(qn('w:val'), "000000")  # Black color
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), "single")
    rPr.append(u)
    new_run.append(rPr)
    
    new_run_text = OxmlElement('w:t')
    new_run_text.text = text
    new_run.append(new_run_text)
    hyperlink.append(new_run)
    
    # Append the hyperlink element to the paragraph.
    paragraph._p.append(hyperlink)
    # Append an empty run to ensure the hyperlink formatting does not continue.
    paragraph.add_run("")
    return hyperlink

def is_email(token):
    """Return True if token is a valid email."""
    token = token.strip()
    return bool(_email_re.match(token))

def is_phone(token):
    """Return True if token consists of (or can be reduced to) exactly 10 digits."""
    digits = ''.join(ch for ch in token if ch.isdigit())
    return len(digits) == 10

def is_url(token):
    """Return True if token is a URL."""
    token = token.strip()
    if _url_re.match(token):
        return True
    if '.' in token and ' ' not in token:
        return True
    return False

class Generator:
    def __init__(self, selected_sections):
        self.selected_sections = selected_sections

    def generate(self, output_file):
        """Generate a Word document resume based on the selected sections."""
        # Create a new document
        doc = Document()
        
        # Set default style to Times New Roman 11pt (Harvard template recommendation)
        style = doc.styles['Normal']
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(11)
        
        # Set document margins (0.5 inches on all sides)
        sections = doc.sections
        for section in sections:
            section.top_margin = Pt(36)    # 0.5 inches = 36 points
            section.bottom_margin = Pt(36)
            section.left_margin = Pt(36)
            section.right_margin = Pt(36)

        # Process each section
        for section in self.selected_sections:
            title = section["title"]
            content = section["content"]

            # Special handling for Personal Information section
            if title == "Personal Information":
                # First item is the name - make it larger, bold, and left-aligned (Harvard template)
                name_paragraph = doc.add_paragraph()
                name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Left-align per Harvard template
                name_run = name_paragraph.add_run(content[0])
                name_run.bold = True
                name_run.font.size = Pt(16)  # 16-18pt per Harvard template
                set_single_spacing(name_paragraph)

                # Remaining items go on one line, separated by bullet points (Harvard template)
                if len(content) > 1:
                    info_paragraph = doc.add_paragraph()
                    info_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Left-align per Harvard template
                    set_single_spacing(info_paragraph)

                    for i, item in enumerate(content[1:], 1):
                        if i > 1:  # Add separator between items (not before first item)
                            separator = info_paragraph.add_run(" • ")  # Bullet point separator per Harvard template
                            separator.font.size = Pt(11)  # Consistent font size

                        # Check if the item is a special token (email, phone, URL)
                        if is_email(item):
                            add_hyperlink(info_paragraph, f"mailto:{item}", item)
                        elif is_phone(item):
                            # Format phone number consistently
                            digits = ''.join(filter(str.isdigit, item))
                            formatted = f"({digits[:3]}) {digits[3:6]}-{digits[6:]}"
                            add_hyperlink(info_paragraph, f"tel:{digits}", formatted)
                        elif is_url(item):
                            url = item if item.startswith('http') else f'http://{item}'
                            add_hyperlink(info_paragraph, url, item)
                        else:
                            run = info_paragraph.add_run(item)
                            run.font.size = Pt(11)  # Consistent font size

                # Add space after personal info with appropriate spacing
                doc.add_paragraph().paragraph_format.space_after = Pt(6)

            # Special handling for Objective section
            elif title == "Objective":
                # Add section header with Harvard template formatting
                header = doc.add_paragraph()
                header.paragraph_format.space_before = Pt(12)  # Space before heading
                header.paragraph_format.space_after = Pt(6)   # Space after heading
                header_run = header.add_run(title.upper())  # Uppercase for section headers
                header_run.bold = True
                header_run.font.size = Pt(12)  # Consistent 12pt for all headers
                
                # Add the selected objective
                obj_para = doc.add_paragraph()
                obj_para.add_run(content)
                set_single_spacing(obj_para)
                # Add space after objective with appropriate spacing
                doc.add_paragraph().paragraph_format.space_after = Pt(6)

            # Special handling for Education and Professional Experience sections
            elif title in ["Education", "Professional Experience"]:
                # Add section header with Harvard template formatting
                header = doc.add_paragraph()
                header.paragraph_format.space_before = Pt(12)  # Space before heading
                header.paragraph_format.space_after = Pt(6)   # Space after heading
                header_run = header.add_run(title.upper())  # Uppercase for section headers
                header_run.bold = True
                header_run.font.size = Pt(12)  # Consistent 12pt for all headers

                # Process each item
                for item in content:
                    # Handle Education items (which are lists) differently
                    if title == "Education" and isinstance(item, list):
                        # Create a paragraph with tab stops for two-column format
                        entry_para = doc.add_paragraph()
                        entry_para.paragraph_format.space_after = Pt(0)
                        entry_para.paragraph_format.space_before = Pt(6)  # Space before each entry
                        
                        # Add a right-aligned tab stop for dates
                        tab_stops = entry_para.paragraph_format.tab_stops
                        tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
                        
                        # Add the institution/degree name (left column)
                        title_run = entry_para.add_run(item[0])
                        title_run.bold = True
                        
                        # If there's a date in the first detail, extract and place it on the right
                        date_text = ""
                        details_to_process = item[1:]
                        
                        if len(details_to_process) > 0 and (":" in details_to_process[0] and 
                                                           ("date" in details_to_process[0].lower() or 
                                                            "graduated" in details_to_process[0].lower())):
                            parts = details_to_process[0].split(":", 1)
                            if len(parts) > 1:
                                date_text = parts[1].strip()
                                # Remove this detail from the list since we've extracted the date
                                details_to_process = details_to_process[1:]
                        
                        # Add the date on the right if we found one
                        if date_text:
                            entry_para.add_run("\t" + date_text)
                        
                        # Remaining items are details as bullet points
                        for detail in details_to_process:
                            detail_para = doc.add_paragraph()
                            detail_para.paragraph_format.left_indent = Pt(18)  # 0.25 inches
                            detail_para.paragraph_format.space_after = Pt(0)
                            detail_run = detail_para.add_run("• " + detail)
                            set_single_spacing(detail_para)
                    else:
                        # Professional Experience items are dictionaries
                        # Create a paragraph with tab stops for two-column format
                        entry_para = doc.add_paragraph()
                        entry_para.paragraph_format.space_after = Pt(0)
                        entry_para.paragraph_format.space_before = Pt(6)  # Space before each entry
                        
                        # Add a right-aligned tab stop for dates
                        tab_stops = entry_para.paragraph_format.tab_stops
                        tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
                        
                        # Add the job title/institution (left column)
                        title_run = entry_para.add_run(item["subtitle"])
                        title_run.bold = True
                        
                        # Add dates if present on the same line (right column)
                        if "date" in item:
                            entry_para.add_run("\t" + item["date"])
                        
                        # Add details as bullet points
                        for detail in item["details"]:
                            detail_para = doc.add_paragraph()
                            detail_para.paragraph_format.left_indent = Pt(18)  # 0.25 inches
                            detail_para.paragraph_format.space_after = Pt(0)
                            detail_run = detail_para.add_run("• " + detail)
                            set_single_spacing(detail_para)

                # Add space after section with appropriate spacing
                doc.add_paragraph().paragraph_format.space_after = Pt(6)

            # Special handling for Core Competencies section
            elif title == "Core Competencies":
                # Add section header with Harvard template formatting
                header = doc.add_paragraph()
                header.paragraph_format.space_before = Pt(12)  # Space before heading
                header.paragraph_format.space_after = Pt(6)   # Space after heading
                header_run = header.add_run(title.upper())  # Uppercase for section headers
                header_run.bold = True
                header_run.font.size = Pt(12)  # Consistent 12pt for all headers

                # Add content items as a comma-separated list (as per user preference)
                if content:
                    comp_para = doc.add_paragraph()
                    comp_para.paragraph_format.left_indent = Pt(18)  # 0.25 inches
                    comp_para.paragraph_format.space_after = Pt(0)
                    comp_para.add_run(", ".join(content))
                    set_single_spacing(comp_para)

            # Default handling for other sections
            else:
                # Add section header with Harvard template formatting
                header = doc.add_paragraph()
                header.paragraph_format.space_before = Pt(12)  # Space before heading
                header.paragraph_format.space_after = Pt(6)   # Space after heading
                header_run = header.add_run(title.upper())  # Uppercase for section headers
                header_run.bold = True
                header_run.font.size = Pt(12)  # Consistent 12pt for all headers

                # Add content items as bullet points
                for item in content:
                    # Check if item is a dictionary with subtitle and date (for projects)
                    if isinstance(item, dict) and "subtitle" in item:
                        # Create a paragraph with tab stops for two-column format
                        entry_para = doc.add_paragraph()
                        entry_para.paragraph_format.space_after = Pt(0)
                        entry_para.paragraph_format.space_before = Pt(6)  # Space before each entry
                        
                        # Add a right-aligned tab stop for dates
                        tab_stops = entry_para.paragraph_format.tab_stops
                        tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
                        
                        # Add the project name (left column)
                        title_run = entry_para.add_run(item["subtitle"])
                        title_run.bold = True
                        
                        # Add dates if present on the same line (right column)
                        if "date" in item:
                            entry_para.add_run("\t" + item["date"])
                        
                        # Add details as bullet points
                        if "details" in item:
                            for detail in item["details"]:
                                detail_para = doc.add_paragraph()
                                detail_para.paragraph_format.left_indent = Pt(18)  # 0.25 inches
                                detail_para.paragraph_format.space_after = Pt(0)
                                detail_run = detail_para.add_run("• " + detail)
                                set_single_spacing(detail_para)
                    else:
                        # Simple string items
                        item_para = doc.add_paragraph()
                        item_para.paragraph_format.left_indent = Pt(18)  # 0.25 inches
                        item_para.paragraph_format.space_after = Pt(0)
                        item_para.add_run("• " + item)
                        set_single_spacing(item_para)

                # Add space after section with appropriate spacing
                doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # Ensure output directory exists
        output_dir = os.path.dirname(output_file)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # Save the document
        doc.save(output_file)
# Example usage (uncomment the lines below to test the function):
# if __name__ == "__main__":
#     import json
#     with open("data.json", "r") as f:
#         selected_sections = json.load(f)
#     generate_resume(selected_sections)
